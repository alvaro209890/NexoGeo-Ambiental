# -*- coding: utf-8 -*-
"""Automação: CONTEXTO AMBIENTAL — Cadastro Ambiental Rural (CAR) -> .docx.

Para cada fazenda cruza:
  - shape do CAR (nº estadual via .dbf)
  - SEMA-MT WFS (situação atual + proprietários)        [requer authkey nos segredos]
  - recibo PDF do CAR (Dados das Áreas dos Imóveis Rurais)
  - APFs_por_CAR.xlsx (APF com situação REGULAR)

Degrada com elegância: sem internet/authkey, usa os proprietários do .dbf e
marca a situação como "—". Porta genérica do antigo ``gerar_contexto_ambiental.py``.

Uso (a partir de software/):
    python -m automations.ctx_ambiental <projeto.json> [saida.docx]
"""
from __future__ import annotations

import os
import re

import openpyxl
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core import geo, io, recibo, secrets
from core import normalize as nz
from core import docx_builder as db
from core.clients import sema

NOME_PADRAO = "Contexto_Ambiental_CAR.docx"
APF_XLSX = "APFs_por_CAR.xlsx"

SIT_MAP = {
    "AGUARDANDO_ANALISE": "Aguardando análise",
    "AGUARDANDO_ENVIO_PRA": "Aguardando análise",
    "EM_ANALISE": "Em análise",
    "CAR_VALIDADO": "Validado (analisado)",
    "ANALISADO": "Analisado",
}


def _apf_regular(xlsx_path: str, nome_fazenda: str):
    """Primeira APF REGULAR da fazenda na planilha, ou ``None``."""
    if not os.path.exists(xlsx_path):
        return None
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    if "APFs por CAR" not in wb.sheetnames:
        return None
    rows = list(wb["APFs por CAR"].iter_rows(values_only=True))
    if len(rows) < 2:
        return None
    ix = {h: i for i, h in enumerate(rows[1])}
    for row in rows[2:]:
        if not row or row[0] is None:
            continue
        if (str(row[ix["CAR / Fazenda"]]).strip() == nome_fazenda
                and str(row[ix["Situação"]]).strip().upper() == "REGULAR"):
            return (row[ix["Nº APF"]], row[ix["Data Validade"]])
    return None


def _coletar(projeto):
    sec = secrets.load_secrets(projeto)
    authkey = sec.get("sema_authkey")
    cars = {c.id: c for c in geo.carregar_cars(projeto)}
    apf_xlsx = os.path.join(projeto.caminho("resultados"), APF_XLSX)

    dados = []
    for fz in projeto.fazendas:
        car = cars[fz.id]
        numest = car.attr(projeto, "num_estadual")
        props = {}
        if authkey and numest:
            try:
                props = sema.consulta_car(numest, authkey) or {}
            except Exception as e:  # serviço fora do ar / sem rede -> degrada
                print(f"  !! SEMA indisponível para {fz.nome}: {e}")
        sit_raw = re.sub(r"[\[\]]", "", str(props.get("SITUACAO", ""))).strip()
        situacao = SIT_MAP.get(sit_raw, sit_raw.replace("_", " ").capitalize()) if sit_raw else "—"
        propr = nz.proprietarios(props.get("NOMESPROPRIETARIOS")
                                 or car.attr(projeto, "proprietarios") or "")
        dados.append({
            "nome": fz.nome, "car": numest, "prop": propr, "sit": situacao,
            "apf": _apf_regular(apf_xlsx, fz.nome),
            "areas": recibo.areas_do_recibo(projeto, fz),
        })
    return dados


def _add_areas(doc, areas, num_id):
    """Insere os documentos do recibo como sub-tópicos (mesmo estilo de texto)."""
    cab = doc.add_paragraph()
    db.set_bullet(cab, 1, num_id)
    cab.paragraph_format.space_before = Pt(4)
    cab.paragraph_format.space_after = Pt(0)
    cab.add_run("Dados das Áreas dos Imóveis Rurais:").font.bold = True
    for documento, tipo, area in areas:
        it = doc.add_paragraph()
        db.set_bullet(it, 2, num_id)
        it.paragraph_format.space_after = Pt(0)
        it.add_run(f"{documento} – {tipo} – {area} ha")


def gerar(projeto, destino: str | None = None) -> str:
    dados = _coletar(projeto)

    doc = db.novo_documento()
    num_id = db.setup_bullets(doc)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rh = h.add_run("CONTEXTO AMBIENTAL")
    rh.font.name = "Tahoma"
    rh.font.bold = True
    rh.font.size = Pt(14)
    rh.font.color.rgb = db.AZUL
    db.add_bottom_border(h)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(10)
    t.paragraph_format.space_after = Pt(8)
    rt = t.add_run("CADASTRO AMBIENTAL RURAL (CAR)")
    rt.font.bold = True
    rt.font.size = Pt(13)
    rt.font.color.rgb = db.PRETO

    for f in dados:
        p = doc.add_paragraph()
        db.set_bullet(p, 0, num_id)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        p.add_run(f["nome"]).font.bold = True
        resto = f" - {f['car']}" if f["car"] else ""
        if f["prop"]:
            resto += f" - {f['prop']}"
        if resto:
            p.add_run(resto)

        if f["areas"]:
            _add_areas(doc, f["areas"], num_id)

        s = doc.add_paragraph()
        db.set_bullet(s, 1, num_id)
        s.paragraph_format.space_after = Pt(0)
        s.add_run("Situação: ").font.bold = True
        s.add_run(f"Ativo – {f['sit']}.")

        a = doc.add_paragraph()
        db.set_bullet(a, 1, num_id)
        a.paragraph_format.space_after = Pt(0)
        a.add_run("APF: ").font.bold = True
        if f["apf"]:
            a.add_run(f"nº {f['apf'][0]} – Regular")
        else:
            a.add_run("Não possui.").font.bold = True

    nota = doc.add_paragraph()
    nota.paragraph_format.space_before = Pt(16)
    rn = nota.add_run(
        "Fonte: situação do CAR extraída do SIMCAR/SEMA-MT "
        "(camada Geoportal:MVW_REQUERIMENTO_ATP) e APFs da consulta pública APF Rural – "
        f"SEMA-MT, em {projeto.data_consulta_efetiva()}."
    )
    rn.font.size = Pt(8)
    rn.font.italic = True
    rn.font.color.rgb = db.CINZA

    destino = destino or io.caminho_resultado(projeto, NOME_PADRAO)
    doc.save(destino)
    return destino


def _main(argv: list[str]) -> int:
    from core.config import load_projeto
    if len(argv) < 2:
        print("uso: python -m automations.ctx_ambiental <projeto.json> [saida.docx]")
        return 2
    proj = load_projeto(argv[1])
    out = gerar(proj, argv[2] if len(argv) > 2 else None)
    print("Gerado:", out)
    return 0


if __name__ == "__main__":
    import sys
    raise SystemExit(_main(sys.argv))
