# -*- coding: utf-8 -*-
"""Automação: CONTEXTO FUNDIÁRIO (certificações INCRA) -> .docx.

Cruza o perímetro de cada fazenda (CAR) com as parcelas certificadas do INCRA
(SIGEF e SNCI), por interseção > 50%. Porta genérica do antigo
``gerar_contexto_fundiario.py`` (lê fazendas/CRS do projeto).

Uso (a partir de software/):
    python -m automations.ctx_fundiario <projeto.json> [saida.docx]
"""
from __future__ import annotations

import os

from shapely.ops import unary_union
from docx.shared import Pt

from core import geo, io
from core import normalize as nz
from core import docx_builder as db
from core.clients import incra

NOME_PADRAO = "Contexto_Fundiario_INCRA.docx"


def _coletar(projeto):
    """Baixa e atribui as certificações por fazenda. Retorna (por_fazenda, totais)."""
    cardir = projeto.caminho("car")
    farm_geoms = {}
    for fz in projeto.fazendas:
        shp = os.path.join(cardir, fz.shape_car, "CAR_ATP.shp")
        farm_geoms[fz.id] = geo.ler_geometria(shp, projeto.crs_geo, unir=True)  # lon/lat (4674)
    validos = [g for g in farm_geoms.values() if g is not None]
    bounds = unary_union(validos).bounds

    brutos = []
    for tema, sistema in ((incra.SIGEF, "SIGEF"), (incra.SNCI, "SNCI")):
        for g, props in incra.parcelas(tema, bounds):
            if g:
                brutos.append((sistema, g, props))

    por = {fz.id: [] for fz in projeto.fazendas}
    for sistema, g, props in brutos:
        melhor, frac_m = None, 0.0
        for fz in projeto.fazendas:
            fg = farm_geoms[fz.id]
            if fg is None:
                continue
            inter = g.intersection(fg)
            if not inter.is_empty:
                frac = inter.area / g.area
                if frac > frac_m:
                    frac_m, melhor = frac, fz.id
        if melhor is None or frac_m < 0.5:
            continue
        if sistema == "SIGEF":
            reg = props.get("registro_matricula")
            cert = {
                "sistema": "SIGEF", "codigo": props.get("parcela_codigo"),
                "data_label": "Data de Aprovação", "data": nz.br_data(props.get("data_aprovacao")),
                "area": geo.reprojetar(g, 4326, projeto.crs_utm).area / 10000.0,
                "situacao": ("SIGEF averbado na matrícula." if reg
                             else "SIGEF não averbado na matrícula."),
            }
        else:
            cert = {
                "sistema": "SNCI", "codigo": props.get("num_certificacao"),
                "data_label": "Data de Certificação", "data": nz.br_data(props.get("data_certificacao")),
                "area": float(props.get("qtd_area_peca_tecnica") or 0),
                "situacao": "SNCI não averbado na matrícula.",
            }
        por[melhor].append(cert)

    for fid in por:
        por[fid].sort(key=lambda c: (c["sistema"] != "SIGEF", c["data"]))
    total = sum(len(v) for v in por.values())
    return por, total


def gerar(projeto, destino: str | None = None) -> str:
    por, total = _coletar(projeto)

    doc = db.novo_documento()
    num_id = db.setup_bullets(doc)

    h = doc.add_paragraph()
    rh = h.add_run("CONTEXTO FUNDIÁRIO")
    rh.font.name = "Tahoma"
    rh.font.bold = True
    rh.font.size = Pt(14)
    rh.font.color.rgb = db.PRETO
    db.add_bottom_border(h)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(8)
    intro.add_run("Possui ")
    intro.add_run(str(total)).font.bold = True
    intro.add_run(" certificações no INCRA (SIGEF e SNCI):")

    for fz in projeto.fazendas:
        certs = por[fz.id]
        if not certs:
            continue
        area_tot = sum(c["area"] for c in certs)
        p = doc.add_paragraph()
        db.set_bullet(p, 0, num_id)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        p.add_run(fz.nome).font.bold = True
        p.add_run(f" com {nz.br(area_tot)} ha")
        for c in certs:
            sc = doc.add_paragraph()
            db.set_bullet(sc, 1, num_id)
            sc.paragraph_format.space_after = Pt(0)
            sc.add_run(f"Certificação ({c['sistema']}): ").font.bold = True
            sc.add_run(f"{c['codigo']} - {c['data_label']}: {c['data']}")
            ss = doc.add_paragraph()
            db.set_bullet(ss, 1, num_id)
            ss.paragraph_format.space_after = Pt(0)
            ss.add_run("Situação: ").font.bold = True
            ss.add_run(c["situacao"])

    m = projeto.municipio
    nota = doc.add_paragraph()
    nota.paragraph_format.space_before = Pt(16)
    rn = nota.add_run(
        "Fonte: Acervo Fundiário INCRA — camadas SIGEF (certificada_sigef_particular_mt) e "
        "SNCI (imoveiscertificados_privado_mt), por interseção com o perímetro de cada imóvel. "
        f"Áreas SIGEF calculadas da geometria certificada (EPSG:{projeto.crs_utm}). "
        f"Consulta em {projeto.data_consulta_efetiva()}."
    )
    rn.font.size = Pt(8)
    rn.font.italic = True
    rn.font.color.rgb = db.CINZA

    destino = destino or io.caminho_resultado(projeto, NOME_PADRAO)
    doc.save(destino)
    return destino


def _main(argv: list[str]) -> int:
    from core.config import load_projeto
    if len(argv) < 2:
        print("uso: python -m automations.ctx_fundiario <projeto.json> [saida.docx]")
        return 2
    proj = load_projeto(argv[1])
    out = gerar(proj, argv[2] if len(argv) > 2 else None)
    print("Gerado:", out)
    return 0


if __name__ == "__main__":
    import sys
    raise SystemExit(_main(sys.argv))
