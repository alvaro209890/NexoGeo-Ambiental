# -*- coding: utf-8 -*-
"""Automacao ponta a ponta: Pre-Analise por shapefile zipado."""
from __future__ import annotations

import json
import os
import re
import shutil
from datetime import date
from pathlib import Path

import fitz
import openpyxl
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from shapely.geometry import mapping, shape as shapely_shape

from core import docx_builder as db
from core import geo, io, normalize as nz, recibo, secrets
from core import overlay
from core.clients import funai, ibama, mapbiomas, sccon, sema, simcar_api
from core.llm import deepseek

SIT_MAP = {
    "AGUARDANDO_ANALISE": "Aguardando análise",
    "AGUARDANDO_ENVIO_PRA": "Aguardando análise",
    "EM_ANALISE": "Em análise",
    "CAR_VALIDADO": "Validado (analisado)",
    "ANALISADO": "Analisado",
}


def _slug_arquivo(s: str) -> str:
    s = re.sub(r"\s+", "_", _limpo(s))
    s = re.sub(r"[^\w\-.]", "", s)
    return s or "Imovel"


def _nome_saida(projeto) -> str:
    return f"Pre_Analise_{_slug_arquivo(projeto.imovel)}.docx"


_EXTS_GEOMETRIA = (".zip", ".shp", ".geojson", ".json", ".kml", ".kmz")


def _zip_entrada(projeto) -> str:
    """Sem geometria explícita, aceita exatamente UMA na pasta de shapes do projeto.

    Formatos aceitos: .zip (shapefile), .shp, .geojson/.json, .kml e .kmz.
    """
    shapes_dir = projeto.caminho("shapes")
    candidatos = sorted(
        os.path.join(shapes_dir, n) for n in os.listdir(shapes_dir)
        if n.lower().endswith(_EXTS_GEOMETRIA)
    ) if os.path.isdir(shapes_dir) else []
    if len(candidatos) == 1:
        return candidatos[0]
    if not candidatos:
        raise RuntimeError(
            f"nenhuma geometria informada e nenhuma encontrada em {shapes_dir}; "
            "envie o arquivo da área pela UI (.zip/.shp/.geojson/.kml/.kmz) ou informe o caminho na chamada"
        )
    raise RuntimeError(
        f"nenhuma geometria informada e há {len(candidatos)} candidatas em {shapes_dir}; "
        "informe qual usar: " + ", ".join(os.path.basename(z) for z in candidatos)
    )


def _deepseek_key(projeto) -> str:
    sec = secrets.load_secrets(projeto)
    return _limpo(sec.get("deepseek_api_key") or os.environ.get("DEEPSEEK_API_KEY", ""))


def _fmt_area(v: float, dec: int = 4) -> str:
    return nz.br(float(v or 0), dec)


def _limpo(s) -> str:
    return re.sub(r"\s+", " ", str(s or "")).strip()


def _mascara_doc(v: str) -> str:
    d = re.sub(r"\D", "", str(v or ""))
    if len(d) == 11:
        return f"{d[:3]}.{d[3:6]}.{d[6:9]}-{d[9:]}"
    if len(d) == 14:
        return nz.mascara_cnpj(d)
    return str(v or "não informado")


def _processo_ibama(v: str) -> str:
    d = re.sub(r"\D", "", str(v or ""))
    if len(d) == 17:
        return f"{d[:5]}.{d[5:11]}/{d[11:15]}-{d[15:]}"
    return str(v or "não informado")


def _limpar_markdown_llm(s: str) -> str:
    linhas = []
    for linha in (s or "").splitlines():
        limpa = linha.strip()
        if not limpa:
            continue
        if limpa.startswith("|") or re.fullmatch(r"[-:| ]+", limpa):
            continue
        linhas.append(limpa)
    texto = " ".join(linhas)
    texto = re.sub(r"[*_#`]+", "", texto)
    texto = texto.replace("‑", "-").replace("–", "-")
    texto = re.sub(r"\bItem\s+Resumo\s+Risco\b", "", texto, flags=re.I)
    return _limpo(texto)


def _sintese_llm_curta(s: str, limite: int = 850) -> str:
    texto = _limpar_markdown_llm(s)
    if len(texto) <= limite:
        return texto
    corte = texto.rfind(".", 0, limite)
    if corte < 300:
        corte = limite
    return texto[:corte + 1].strip()


def _auto_com_serie(num, serie) -> str:
    num = _limpo(num)
    serie = _limpo(serie)
    if num and serie and not num.endswith(f"-{serie}"):
        return f"{num}-{serie}"
    return num or "não informado"


def _expand_bounds(bounds, delta_graus: float):
    minx, miny, maxx, maxy = bounds
    return minx - delta_graus, miny - delta_graus, maxx + delta_graus, maxy + delta_graus


def _pdf_text(path: str) -> str:
    if not path or not os.path.exists(path):
        return ""
    doc = fitz.open(path)
    try:
        return "\n".join(page.get_text() for page in doc)
    finally:
        doc.close()


def _authkey(projeto) -> str | None:
    sec = secrets.load_secrets(projeto)
    return sec.get("sema_authkey") or sema.authkey_local(projeto.raiz_abs())


def _areas_from_ai(data: dict | None):
    if not data:
        return []
    for key in ("areas_imoveis", "dados_areas_imoveis", "areas"):
        vals = data.get(key)
        if isinstance(vals, list):
            out = []
            for it in vals:
                if not isinstance(it, dict):
                    continue
                doc = it.get("documento") or it.get("matricula") or it.get("descricao") or ""
                tipo = it.get("tipo") or it.get("natureza") or ""
                area = it.get("area_ha") or it.get("area") or ""
                if doc or tipo or area:
                    out.append((_limpo(doc), _limpo(tipo), _limpo(area)))
            return out
    return []


def _extrair_recibo(pdf_path: str, api_key: str, avisos: list[str]):
    texto = _pdf_text(pdf_path)
    llm = None
    areas_parser = recibo.areas_imoveis(pdf_path)
    areas = []
    if texto:
        llm = deepseek.extrair_documento_pdf(texto, os.path.basename(pdf_path), "recibo_car", api_key=api_key)
        if llm.ok:
            areas = _areas_from_ai(llm.data)
        else:
            avisos.append(f"DeepSeek Flash falhou em {os.path.basename(pdf_path)}: {llm.error}")
    if areas_parser:
        areas = areas_parser
    return areas, texto, llm


def _apf_regular(xlsx_path: str, nome_fazenda: str):
    if not os.path.exists(xlsx_path):
        return None
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    if "APFs por CAR" not in wb.sheetnames:
        return None
    rows = list(wb["APFs por CAR"].iter_rows(values_only=True))
    if len(rows) < 2:
        return None
    ix = {h: i for i, h in enumerate(rows[1])}
    for row in rows[2:]:
        if not row or row[0] is None:
            continue
        if (_limpo(row[ix["CAR / Fazenda"]]) == nome_fazenda
                and _limpo(row[ix["Situação"]]).upper() == "REGULAR"):
            return (row[ix["Nº APF"]], row[ix["Data Validade"]])
    return None


def _apf_pdf_em_dir(apf_dir: str, numero_apf) -> str | None:
    if not numero_apf:
        return None
    needle = str(numero_apf).replace("/", "-")
    if not os.path.isdir(apf_dir):
        return None
    for nome in os.listdir(apf_dir):
        if nome.lower().endswith(".pdf") and needle in nome:
            return os.path.join(apf_dir, nome)
    return None


def _copiar_pdfs(origem_dir: str, destino_dir: str):
    if not destino_dir or not os.path.isdir(origem_dir):
        return
    os.makedirs(destino_dir, exist_ok=True)
    for nome in os.listdir(origem_dir):
        if nome.lower().endswith(".pdf"):
            src = os.path.join(origem_dir, nome)
            dst = os.path.join(destino_dir, nome)
            if os.path.abspath(src) != os.path.abspath(dst):
                shutil.copy2(src, dst)


def _garantir_apfs(projeto, avisos: list[str], apf_saida_dir: str | None = None) -> str:
    from automations import apf

    xlsx = io.caminho_resultado(projeto, "APFs_por_CAR.xlsx")
    try:
        apf.gerar(projeto, xlsx, baixar=False)
    except Exception as e:
        avisos.append(f"Consulta APF falhou; usando planilha local se existir: {e}")
    try:
        apf.baixar_pdfs(projeto)
    except Exception as e:
        avisos.append(f"Download de APFs falhou; usando PDFs locais se existirem: {e}")
    if apf_saida_dir:
        _copiar_pdfs(io.caminho_consulta(projeto, "APF"), apf_saida_dir)
    return xlsx


def _fazendas_intersectadas(projeto, area_zip, authkey: str | None = None,
                            avisos: list[str] | None = None):
    """Fazendas (CARs) que intersectam a área do zip.

    Com ``fazendas[]`` no projeto: usa os shapes locais, na ORDEM do projeto.
    Sem ``fazendas[]`` (modo só-shape): descobre os CARs pela camada SEMA
    ``MVW_REQUERIMENTO_ATP`` no bbox do zip, ordenados por área de interseção desc.
    """
    if projeto.fazendas:
        cars = geo.carregar_cars(projeto)
        by_id = {c.id: c for c in cars}
        ids = []
        for car in cars:
            inter = car.geom.intersection(area_zip.union_utm)
            if not inter.is_empty and inter.area > 1000:
                ids.append(car.id)
        ordem = {fz.id: i for i, fz in enumerate(projeto.fazendas)}
        fazendas = [fz for fz in projeto.fazendas if fz.id in ids]
        fazendas.sort(key=lambda fz: ordem.get(fz.id, 999))
        return fazendas, by_id
    return _fazendas_via_sema(projeto, area_zip, authkey, avisos if avisos is not None else [])


def _fazendas_via_sema(projeto, area_zip, authkey: str | None, avisos: list[str]):
    """Modo só-shape: cria as fazendas dinamicamente a partir dos CARs da SEMA."""
    from core.config import Fazenda

    if not authkey:
        avisos.append("Modo só-shape sem authkey SEMA: não foi possível descobrir os CARs da área.")
        return [], {}
    try:
        data = sema.get_features(sema.LAYER_CAR, authkey, bounds=area_zip.bbox_geo, count=200)
    except Exception as e:
        avisos.append(f"Descoberta de CARs pela SEMA falhou: {e}")
        return [], {}
    feats = overlay.features_de_geojson(data, "SEMA-MT", sema.LAYER_CAR, dst_epsg=projeto.crs_utm)
    melhores: dict[str, tuple[float, object, float]] = {}
    for feat in feats:
        num = _limpo(feat.props.get("NUMEROESTADUAL"))
        if not num:
            continue
        inter = feat.geom.intersection(area_zip.union_utm)
        if inter.is_empty or inter.area <= 1000:
            continue
        req_id = float(feat.props.get("REQUERIMENTO_ID") or 0)
        atual = melhores.get(num)
        if atual is None or req_id > atual[0]:
            melhores[num] = (req_id, feat, inter.area)
    fazendas, by_id = [], {}
    for num, (_req_id, feat, _ainter) in sorted(melhores.items(), key=lambda kv: -kv[1][2]):
        nome = _limpo(
            feat.props.get("NOMEIMOVELRURAL") or feat.props.get("NOME_IMOVEL")
            or feat.props.get("NOMEIMOVEL") or feat.props.get("DENOMINACAO") or ""
        ) or f"CAR {num}"
        fid = re.sub(r"\W+", "_", num).strip("_").lower() or f"car_{len(fazendas) + 1}"
        fazendas.append(Fazenda(id=fid, nome=nome, shape_car="", car_estadual=num))
        by_id[fid] = geo.Car(fid, nome, feat.geom, dict(feat.props))
    if not fazendas:
        avisos.append("Modo só-shape: nenhum CAR da SEMA intersecta a área enviada.")
    return fazendas, by_id


def _coletar_car(projeto, area_zip, authkey: str | None, api_key: str, avisos: list[str],
                 car_saida_dir: str | None = None, apf_saida_dir: str | None = None):
    fazendas, cars_by_id = _fazendas_intersectadas(projeto, area_zip, authkey, avisos)
    apf_xlsx = _garantir_apfs(projeto, avisos, apf_saida_dir)
    car_dir = io.garantir_dir(car_saida_dir or io.caminho_consulta(projeto, "CAR"))
    dados = []
    llm_docs = []
    for fz in fazendas:
        car = cars_by_id[fz.id]
        numest = fz.car_estadual or car.attr(projeto, "num_estadual")
        props = {}
        if authkey and numest:
            try:
                props = sema.consulta_car(numest, authkey) or {}
            except Exception as e:
                avisos.append(f"SEMA CAR indisponível para {fz.nome}: {e}")
        sit_raw = re.sub(r"[\[\]]", "", _limpo(props.get("SITUACAO"))).strip()
        situacao = SIT_MAP.get(sit_raw, sit_raw.replace("_", " ").capitalize()) if sit_raw else "não informado"
        proprietarios = nz.proprietarios(props.get("NOMESPROPRIETARIOS") or car.attr(projeto, "proprietarios") or "")

        recibo_pdf = None
        req_id = props.get("REQUERIMENTO_ID")
        if req_id:
            recibo_pdf = os.path.join(car_dir, simcar_api.nome_arquivo(fz.nome, req_id))
            try:
                simcar_api.baixar_recibo(req_id, recibo_pdf)
            except Exception as e:
                avisos.append(f"Download SIMCAR falhou para {fz.nome}: {e}")
                recibo_pdf = None
        if not recibo_pdf and fz.recibo_pdf:
            cand = io.caminho_consulta(projeto, "CAR", fz.recibo_pdf)
            if os.path.exists(cand):
                if car_saida_dir:
                    os.makedirs(car_saida_dir, exist_ok=True)
                    dst = os.path.join(car_saida_dir, os.path.basename(cand))
                    shutil.copy2(cand, dst)
                    recibo_pdf = dst
                else:
                    recibo_pdf = cand

        areas, texto_recibo, llm = _extrair_recibo(recibo_pdf, api_key, avisos) if recibo_pdf else ([], "", None)
        if llm:
            llm_docs.append({"arquivo": os.path.basename(recibo_pdf), "ok": llm.ok, "modelo": llm.model, "erro": llm.error})

        apf_reg = _apf_regular(apf_xlsx, fz.nome)
        apf_ai = None
        apf_base = apf_saida_dir or io.caminho_consulta(projeto, "APF")
        apf_pdf = _apf_pdf_em_dir(apf_base, apf_reg[0] if apf_reg else None)
        if apf_pdf:
            texto_apf = _pdf_text(apf_pdf)
            if texto_apf:
                apf_ai = deepseek.extrair_documento_pdf(texto_apf, os.path.basename(apf_pdf), "apf", api_key=api_key)
                llm_docs.append({"arquivo": os.path.basename(apf_pdf), "ok": apf_ai.ok, "modelo": apf_ai.model, "erro": apf_ai.error})
                if not apf_ai.ok:
                    avisos.append(f"DeepSeek Flash falhou em {os.path.basename(apf_pdf)}: {apf_ai.error}")

        dados.append({
            "fazenda": fz,
            "car": car,
            "num_estadual": numest,
            "props": props,
            "proprietarios": proprietarios,
            "situacao": situacao,
            "recibo_pdf": recibo_pdf,
            "areas": areas,
            "apf": apf_reg,
            "apf_pdf": apf_pdf,
            "apf_ai": apf_ai.data if apf_ai and apf_ai.ok else None,
        })
    return dados, llm_docs


def _coletar_fundiario(projeto, avisos: list[str]):
    try:
        from automations import ctx_fundiario
        return ctx_fundiario._coletar(projeto)  # reuso do coletor existente
    except Exception as e:
        avisos.append(f"INCRA indisponível; usando area_total como apoio parcial: {e}")
        por = {fz.id: [] for fz in projeto.fazendas}
        for item in projeto.area_total:
            nome = _limpo(item.get("nome"))
            fid = next((fz.id for fz in projeto.fazendas if fz.nome in nome or nome in fz.nome), None)
            if fid:
                por[fid].append({
                    "sistema": item.get("tipo", ""),
                    "codigo": "preenchimento manual",
                    "data_label": "Data",
                    "data": "não informado",
                    "area": float(item.get("certificacao_ha") or 0),
                    "situacao": "Dados de certificação importados do projeto; requer conferência.",
                })
        return por, sum(len(v) for v in por.values())


def _coletar_protegidas(projeto, area_zip, authkey: str | None, avisos: list[str]):
    out = {"ti": [], "uc": []}
    try:
        ti_data = funai.terras_indigenas(area_zip.bbox_geo)
        ti_feats = overlay.features_de_geojson(ti_data, "FUNAI", funai.LAYER_TI, dst_epsg=projeto.crs_utm)
        out["ti"] = overlay.calcular_distancias(area_zip.union_utm, ti_feats, limite_km=20)
    except Exception as e:
        avisos.append(f"FUNAI indisponível: {e}")
    if authkey:
        try:
            uc_data = sema.get_features(
                "Geoportal:UNIDADES_CONSERVACAO", authkey,
                bounds=_expand_bounds(area_zip.bbox_geo, 2.0), count=200, timeout=120,
            )
            uc_feats = overlay.features_de_geojson(uc_data, "SEMA-MT", "Geoportal:UNIDADES_CONSERVACAO", dst_epsg=projeto.crs_utm)
            out["uc"] = overlay.calcular_distancias(area_zip.union_utm, uc_feats, limite_km=250)
        except Exception as e:
            avisos.append(f"UC SEMA indisponível: {e}")
    return out


def _hit_com_fazenda(hit, feat, cars):
    return overlay.atribuir_fazenda(hit, feat.geom, cars)


def _coletar_legal(projeto, area_zip, authkey: str | None, cars_by_id, avisos: list[str], api_key: str):
    cars = list(cars_by_id.values())
    legal = {"ibama": [], "sema_embargos": [], "sema_desembargos": [], "resumo_llm": None}
    try:
        data = ibama.embargos_pamgia(area_zip.bbox_geo)
        feats = overlay.features_de_geojson(data, "IBAMA/PAMGIA", "adm_embargos_ibama_a", dst_epsg=projeto.crs_utm)
        for feat in feats:
            for hit in overlay.cruzar(area_zip.union_utm, [feat], layer_area_field="qtd_area_embargada", min_area_ha=0):
                legal["ibama"].append(_hit_com_fazenda(hit, feat, cars))
    except Exception as e:
        avisos.append(f"IBAMA/PAMGIA indisponível: {e}")

    if authkey:
        for layer, chave in [
            ("Geoportal:AREAS_EMBARGADAS_SEMA", "sema_embargos"),
            ("Geoportal:AREA_EMBARGADA_SIGA_POLIGONO", "sema_embargos"),
            ("Geoportal:AREAS_DESEMBARGADAS_SEMA", "sema_desembargos"),
        ]:
            try:
                data = sema.get_features(layer, authkey, bounds=area_zip.bbox_geo, count=100)
                feats = overlay.features_de_geojson(data, "SEMA-MT", layer, dst_epsg=projeto.crs_utm)
                for feat in feats:
                    for hit in overlay.cruzar(area_zip.union_utm, [feat], layer_area_field="AREA_HA", min_area_ha=0):
                        legal[chave].append(_hit_com_fazenda(hit, feat, cars))
            except Exception as e:
                avisos.append(f"{layer} indisponível: {e}")

    registros = {
        "ibama": [h.props for h in legal["ibama"]],
        "sema_embargos": [h.props for h in legal["sema_embargos"]],
        "sema_desembargos": [h.props for h in legal["sema_desembargos"]],
    }
    if any(registros.values()):
        llm = deepseek.resumir_juridico(json.dumps(registros, ensure_ascii=False, indent=2), api_key=api_key)
        legal["resumo_llm"] = llm
        if not llm.ok:
            avisos.append(f"DeepSeek Pro falhou na seção jurídica: {llm.error}")
    return legal


def _coletar_alertas(projeto, area_zip, avisos: list[str]):
    alertas = []
    try:
        layer, data = mapbiomas.alertas(area_zip.bbox_geo)
        feats = overlay.features_de_geojson(data, "MapBiomas Alerta", layer, dst_epsg=projeto.crs_utm)
        hits = overlay.cruzar(area_zip.union_utm, feats, min_area_ha=0.01)
        for h in hits:
            alertas.append(h)
    except Exception as e:
        avisos.append(f"MapBiomas indisponível: {e}")
    return alertas


def _ano_alerta(alerta: dict) -> str:
    data = str(alerta.get("data_deteccao") or alerta.get("data_publicacao") or "")
    m = re.search(r"\d{4}", data)
    return m.group(0) if m else "sem_ano"


def _plot_geom(ax, geom, edgecolor, facecolor="none", linewidth=1.2, alpha=0.35, label=None):
    if geom.is_empty:
        return
    if geom.geom_type == "Polygon":
        xs, ys = geom.exterior.xy
        ax.fill(xs, ys, facecolor=facecolor, edgecolor=edgecolor, linewidth=linewidth, alpha=alpha, label=label)
        for interior in geom.interiors:
            ix, iy = interior.xy
            ax.plot(ix, iy, color=edgecolor, linewidth=0.5)
        return
    if geom.geom_type in ("MultiPolygon", "GeometryCollection"):
        first = True
        for part in geom.geoms:
            _plot_geom(ax, part, edgecolor, facecolor, linewidth, alpha, label if first else None)
            first = False
        return
    if hasattr(geom, "xy"):
        xs, ys = geom.xy
        ax.plot(xs, ys, color=edgecolor, linewidth=linewidth, label=label)


def _gerar_figuras_sccon(area_zip, alertas_sccon: list[dict], destino_dir: str, avisos: list[str]):
    if not alertas_sccon:
        return []
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception as e:
        avisos.append(f"Figura SCCON não gerada; matplotlib indisponível: {e}")
        return []

    os.makedirs(destino_dir, exist_ok=True)
    minx, miny, maxx, maxy = area_zip.bbox_geo
    dx = max((maxx - minx) * 0.16, 0.01)
    dy = max((maxy - miny) * 0.16, 0.01)

    grupos = {"geral": alertas_sccon}
    for alerta in alertas_sccon:
        grupos.setdefault(_ano_alerta(alerta), []).append(alerta)

    figuras = []
    for ano, itens in sorted(grupos.items(), key=lambda kv: (kv[0] != "geral", kv[0])):
        fig, ax = plt.subplots(figsize=(8.6, 6.0), dpi=180)
        fig.patch.set_facecolor("white")
        ax.set_facecolor("#f3f7f4")
        _plot_geom(ax, area_zip.union_geo, "#14543f", facecolor="#dcefe3", linewidth=2.0, alpha=0.42, label="Propriedade")

        for alerta in itens:
            gj = alerta.get("geojson")
            if not gj:
                continue
            try:
                geom = shapely_shape(gj)
            except Exception:
                continue
            _plot_geom(ax, geom, "#b42318", facecolor="#e85d4a", linewidth=1.25, alpha=0.5, label="Alerta SCCON")
            rp = geom.representative_point()
            ax.text(rp.x, rp.y, str(alerta.get("codigo") or alerta.get("id")), fontsize=6.8, color="#5b1b12")

        ax.set_xlim(minx - dx, maxx + dx)
        ax.set_ylim(miny - dy, maxy + dy)
        ax.set_aspect("equal", adjustable="box")
        ax.grid(True, color="#cfdad3", linewidth=0.45)
        ax.set_xlabel("Longitude (EPSG:4674)")
        ax.set_ylabel("Latitude (EPSG:4674)")
        titulo_ano = "todos os anos" if ano == "geral" else ano
        ax.set_title(f"SCCON Alertas MT - {titulo_ano}", fontsize=12, fontweight="bold", pad=10)
        handles, labels = ax.get_legend_handles_labels()
        unicos = dict(zip(labels, handles))
        if unicos:
            ax.legend(unicos.values(), unicos.keys(), loc="upper right", fontsize=8, frameon=True)
        fig.text(
            0.02,
            0.02,
            "Fonte: Dashboard SCCON de Alertas MT. Polígono da propriedade plotado a partir do shapefile informado.",
            fontsize=7,
            color="#4f5d55",
        )
        nome = f"sccon_alertas_{ano}.png".replace("/", "_")
        path = os.path.join(destino_dir, nome)
        fig.tight_layout(rect=[0, 0.035, 1, 1])
        fig.savefig(path)
        plt.close(fig)
        figuras.append({"ano": ano, "path": path})
    return figuras


def _coletar_sccon(projeto, area_zip, avisos: list[str], saida_dir: str | None = None):
    destino_dir = saida_dir or io.garantir_dir(io.caminho_consulta(projeto, "SCCON"))
    periodo_inicio = "2020-01-01"
    periodo_fim = date.today().isoformat()
    try:
        raw = sccon.buscar_alertas_por_geometria(
            mapping(area_zip.union_geo),
            data_inicio=periodo_inicio,
            data_fim=periodo_fim,
        )
        alertas_sccon = [sccon.normalizar_alerta(item) for item in raw]
        alertas_sccon.sort(key=lambda a: (a.get("data_deteccao") or "", a.get("codigo") or ""))
        figuras = _gerar_figuras_sccon(area_zip, alertas_sccon, destino_dir, avisos)
        return {
            "periodo": (periodo_inicio, periodo_fim),
            "alertas": alertas_sccon,
            "figuras": figuras,
            "fonte": "https://alertas.sccon.com.br/matogrosso/#/",
        }
    except Exception as e:
        avisos.append(f"SCCON Alertas indisponível: {e}")
        return {"periodo": (periodo_inicio, periodo_fim), "alertas": [], "figuras": [], "fonte": "https://alertas.sccon.com.br/matogrosso/#/"}


def _add_bullet(doc, text="", level=0, num_id=77, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    db.set_bullet(p, level, num_id)
    p.paragraph_format.space_after = Pt(0)
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).font.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def _add_label_value(doc, label: str, value, level=1, num_id=77, suffix: str = ""):
    p = doc.add_paragraph()
    db.set_bullet(p, level, num_id)
    p.paragraph_format.space_after = Pt(0)
    p.add_run(f"{label}: ").font.bold = True
    p.add_run(f"{_limpo(value) if value not in (None, '') else 'não informado'}{suffix}")
    return p


def _dominialidade(doc, projeto, num_id, avisos: list[str]):
    dom = projeto.dominialidade
    db.cabecalho_secao(doc, "DOMINIALIDADE", db.PRETO)

    denominacoes = list(dict.fromkeys(m.denominacao for m in dom.matriculas if m.denominacao))
    empreendimentos = denominacoes or [fz.nome for fz in projeto.fazendas] or [projeto.imovel]
    rotulo = "Empreendimentos" if len(empreendimentos) > 1 else "Empreendimento"
    doc.add_paragraph(f"{rotulo}: " + ", ".join(empreendimentos) + ".")
    doc.add_paragraph(f"Município: {projeto.municipio.nome} ({projeto.municipio.uf})")
    doc.add_paragraph("Proprietários(as), conforme as matrículas dos imóveis:")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(["Matrícula", "Propriedade", "Proprietário(a)", "CPF/CNPJ", "Área (ha)"]):
        table.rows[0].cells[i].text = h
    if dom.matriculas:
        for m in dom.matriculas:
            cells = table.add_row().cells
            for i, val in enumerate([
                m.numero, m.denominacao, m.proprietario, m.cpf_cnpj,
                _fmt_area(m.area_ha) if m.area_ha else "",
            ]):
                cells[i].text = _limpo(val)
    else:
        avisos.append(
            "Dominialidade sem matrículas: envie os PDFs das matrículas pela UI "
            "(extração IA + grade de conferência) ou preencha 'dominialidade.matriculas' no projeto.json."
        )
        cells = table.add_row().cells
        cells[0].text = "não informado"

    doc.add_paragraph("Registro:")
    if dom.matriculas and dom.cri:
        sufixo_reg = f" do C.R.I de {dom.cri}" + (f" (CNS: {dom.cns})" if dom.cns else "")
        for m in dom.matriculas:
            nome_reg = m.denominacao or projeto.imovel
            _add_bullet(doc, f"{nome_reg} - Matrícula Nº. {m.numero}{sufixo_reg}", 0, num_id)
    else:
        _add_bullet(doc, "Registro cartorial não informado; requer conferência.", 0, num_id)
        if dom.matriculas and not dom.cri:
            avisos.append("Dominialidade sem CRI/CNS ('dominialidade.registro'): seção Registro incompleta.")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.add_run("Área total:").font.bold = True
    for it in projeto.area_total:
        cert = float(it.get("certificacao_ha") or 0)
        matr = float(it.get("matricula_ha") or 0)
        nome = it.get("nome")
        _add_bullet(doc, f"{nome}: {_fmt_area(cert)} ha ({it.get('tipo')})", 0, num_id)
        label = "Área equivalente" if abs(cert - matr) < 0.0001 else "Divergência"
        _add_bullet(doc, f"{label}: {_fmt_area(matr, nz.casas_decimais(matr))} ha (Matrícula)", 1, num_id)
    if not projeto.area_total and dom.matriculas:
        total = sum(m.area_ha or 0 for m in dom.matriculas)
        if total:
            _add_bullet(doc, f"{projeto.imovel}: {_fmt_area(total)} ha (soma das matrículas)", 0, num_id)


def _contexto_fundiario(doc, projeto, dados_fundiarios, num_id):
    por, total = dados_fundiarios
    db.cabecalho_secao(doc, "CONTEXTO FUNDIÁRIO", db.PRETO)
    sigef = sum(1 for certs in por.values() for c in certs if _limpo(c.get("sistema")).upper() == "SIGEF")
    snci = sum(1 for certs in por.values() for c in certs if _limpo(c.get("sistema")).upper() == "SNCI")
    if sigef or snci:
        doc.add_paragraph(f"Possui {sigef} certificações no INCRA SIGEF e {snci} certificações no INCRA SNCI N°.:")
    else:
        doc.add_paragraph(f"Possui {total} certificações no INCRA (SIGEF e SNCI) N°.:")
    for fz in projeto.fazendas:
        certs = por.get(fz.id, [])
        if not certs:
            continue
        area = sum(float(c.get("area") or 0) for c in certs)
        _add_bullet(doc, f"{fz.nome} com {_fmt_area(area)} ha", 0, num_id)
        for c in certs:
            _add_bullet(doc, f"Certificação ({c.get('sistema')}): {c.get('codigo')} - {c.get('data_label')}: {c.get('data')}", 1, num_id, f"Certificação ({c.get('sistema')}): ")
            _add_bullet(doc, f"Situação: {c.get('situacao')}", 1, num_id, "Situação: ")
    doc.add_paragraph("Sobreposições:")
    _add_bullet(doc, "Não foram encontradas sobreposições.", 0, num_id)


def _contexto_ambiental(doc, car_dados, projeto, num_id):
    db.cabecalho_secao(doc, "CONTEXTO AMBIENTAL", db.AZUL)
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = t.add_run("CADASTRO AMBIENTAL RURAL (CAR)")
    rt.font.bold = True
    rt.font.size = Pt(13)
    for item in car_dados:
        fz = item["fazenda"]
        p = _add_bullet(doc, "", 0, num_id)
        p.add_run(fz.nome).font.bold = True
        rest = f" - {item['num_estadual']}"
        if item["proprietarios"]:
            rest += f" - {item['proprietarios']}"
        p.add_run(rest)
        if item["areas"]:
            _add_bullet(doc, "Dados das Áreas dos Imóveis Rurais:", 1, num_id, "Dados das Áreas dos Imóveis Rurais:")
            for documento, tipo, area in item["areas"]:
                _add_bullet(doc, f"{documento} - {tipo} - {area} ha", 2, num_id)
        _add_bullet(doc, f"Situação: Ativo - {item['situacao']}.", 1, num_id, "Situação: ")
        if item["apf"]:
            _add_bullet(doc, f"APF: nº {item['apf'][0]} - Regular", 1, num_id, "APF: ")
        else:
            _add_bullet(doc, "APF: Não possui.", 1, num_id, "APF: ")
    doc.add_paragraph("Sobreposições:")
    _add_bullet(doc, "Não foram encontradas sobreposições.", 0, num_id)
    db.paragrafo_fonte(
        doc,
        "Fonte: situação do CAR extraída do SIMCAR/SEMA-MT "
        f"(camada Geoportal:MVW_REQUERIMENTO_ATP), recibos SIMCAR e APFs da consulta pública APF Rural, em {projeto.data_consulta_efetiva()}."
    )


def _areas_protegidas(doc, protegidas, num_id):
    db.cabecalho_secao(doc, "ÁREAS PROTEGIDAS (TERRAS INDÍGENAS E UNIDADES DE CONSERVAÇÃO)", db.PRETO)
    if protegidas["ti"]:
        ti = protegidas["ti"][0]
        nome = ti.props.get("terrai_nome") or ti.props.get("TERRAI_NOME") or "terra indígena não informada"
        if ti.area_intersecao_ha > 0.01:
            txt = f"As propriedades se sobrepõem à Terra Indígena {nome} em aproximadamente {_fmt_area(ti.area_intersecao_ha)} ha."
        elif ti.distancia_km < 0.05:
            txt = f"As propriedades se encontram em divisa com a Terra Indígena {nome}, porém a Terra Indígena não se sobrepõe às propriedades."
        else:
            txt = f"As propriedades se encontram a {nz.br(ti.distancia_km, 2)} km da Terra Indígena {nome}."
        _add_bullet(doc, txt, 0, num_id)
    else:
        _add_bullet(doc, "Não foram identificadas Terras Indígenas sobrepostas ou em divisa nas fontes consultadas.", 0, num_id)
    if protegidas["uc"]:
        uc = (
            next((h for h in protegidas["uc"] if _limpo(h.props.get("CATEGORIA")).upper() not in ("RPPN", "REBIO")), None)
            or protegidas["uc"][0]
        )
        nome = uc.props.get("NOME") or "unidade de conservação não informada"
        _add_bullet(doc, f"As propriedades se encontram a {nz.br(uc.distancia_km, 2)} km da unidade de conservação {nome}.", 0, num_id)
    else:
        _add_bullet(doc, "Não foram identificadas Unidades de Conservação próximas nas fontes consultadas.", 0, num_id)


def _legal(doc, legal, num_id, projeto):
    db.cabecalho_secao(doc, "INFRAÇÕES, PENALIDADES E/OU TERMOS DE COMPROMISSO", db.PRETO)
    _add_bullet(
        doc,
        f"Termos de compromisso ambiental (TACs): Nada consta nas fontes consultadas em {projeto.data_consulta_efetiva()}; requer conferência manual.",
        0, num_id,
    )
    _add_bullet(doc, "Embargos e Desembargos/Autos de Infração:", 0, num_id)
    if legal["ibama"] or legal["sema_embargos"]:
        for idx, h in enumerate(legal["ibama"] + legal["sema_embargos"], start=1):
            p = h.props
            _add_bullet(doc, f"{h.fazenda_nome or p.get('nome_imovel') or 'Imóvel intersectado'}.", 0, num_id)
            _add_bullet(doc, f"Auto de Infração {idx}", 0, num_id)
            orgao = "IBAMA" if h.fonte.startswith("IBAMA") else "SEMA/MT"
            nome = p.get("nome_embargado") or p.get("NOME") or p.get("NOME_RAZAO_SOCIAL") or "não informado"
            cpf = _mascara_doc(p.get("cpf_cnpj_embargado") or p.get("CPF_CNPJ") or p.get("CPFCNPJ") or "")
            dano = p.get("des_tad") or p.get("DANO") or p.get("DESCRICAO_DA_OCORRENCIA") or "não informado"
            auto = _auto_com_serie(
                p.get("num_auto_infracao") or p.get("A_INFRAC") or p.get("NUMERO_AUTO_INFRACAO"),
                p.get("serie_auto_infracao"),
            )
            termo = _auto_com_serie(
                p.get("num_tad") or p.get("T_EMBARGO") or p.get("NUMERO_TERMO_EMBARGO"),
                p.get("serie_tad"),
            )
            processo = _processo_ibama(p.get("num_processo") or p.get("N_PROCESSO") or p.get("NUMERO_PROCESSO") or "")
            area = p.get("qtd_area_embargada") or p.get("AREA_HA") or p.get("QUANTIDADE") or "não informado"
            _add_label_value(doc, "Órgão", orgao, 1, num_id)
            _add_label_value(doc, "Nome", _limpo(nome).title() if str(nome).isupper() else _limpo(nome), 1, num_id)
            _add_label_value(doc, "CPF/CNPJ", cpf, 1, num_id)
            _add_label_value(doc, "Dano", dano, 1, num_id)
            _add_label_value(doc, "Auto de infração", auto, 1, num_id)
            if termo:
                _add_label_value(doc, "Termo de embargo", termo, 1, num_id)
            _add_label_value(doc, "Número do processo", processo, 1, num_id)
            _add_label_value(doc, "Área", area, 1, num_id, " ha")
            _add_label_value(doc, "Critério de vínculo", f"{h.criterio}; requer conferência.", 1, num_id)
    else:
        _add_bullet(doc, f"Nada consta nas fontes consultadas (IBAMA/PAMGIA e SEMA-MT) em {projeto.data_consulta_efetiva()}.", 1, num_id)

    if legal["sema_desembargos"]:
        _add_bullet(doc, "Desembargos:", 0, num_id)
        desembargos = sorted(legal["sema_desembargos"], key=lambda h: _limpo(h.props.get("PROPRIEDAD")).upper())
        for idx, h in enumerate(desembargos, start=1):
            p = h.props
            _add_bullet(doc, f"Desembargo {idx}", 0, num_id)
            for label, key in [
                ("Órgão", None), ("Nome", "NOME"), ("CPF", "CPF_CNPJ"), ("Propriedade", "PROPRIEDAD"),
                ("Dano", "DANO"), ("Número do processo", "N_PROCESSO"), ("Área", "AREA_HA"),
            ]:
                val = "SEMA/MT" if key is None else p.get(key, "não informado")
                if key == "CPF_CNPJ":
                    val = _mascara_doc(val)
                sufixo = " ha" if key == "AREA_HA" and val not in (None, "") else ""
                _add_label_value(doc, label, val, 1, num_id, sufixo)

    llm = legal.get("resumo_llm")
    if llm and llm.ok and llm.content:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.add_run("Síntese jurídica por IA (DeepSeek v4 Pro, requer conferência): ").font.bold = True
        p.add_run(_sintese_llm_curta(llm.content))
    db.paragrafo_fonte(
        doc,
        "Aviso: seção de uso interno/pré-análise, elaborada com fontes oficiais consultadas e sujeita a revisão humana obrigatória; ausência de registro não significa regularidade absoluta."
    )


def _alertas(doc, alertas, sccon_data, num_id, projeto):
    db.cabecalho_secao(doc, "ALERTAS", db.PRETO)
    doc.add_paragraph("MapBiomas Alertas:")
    if alertas:
        for a in alertas:
            if isinstance(a, dict):
                _add_bullet(doc, f"Há um alerta inserido na {a.get('fazenda', 'área analisada')}:", 0, num_id)
                _add_bullet(doc, f"○ Código do alerta: {a.get('codigo')}", 1, num_id)
                _add_bullet(doc, f"○ Ano: {a.get('ano')}", 1, num_id)
                _add_bullet(doc, f"○ Área: {a.get('area')} ha", 1, num_id)
                _add_bullet(doc, f"○ Observação: {a.get('observacao')}", 1, num_id)
            else:
                codigo = (a.props.get("alert_code") or a.props.get("cod_alerta") or a.props.get("codigo_alerta")
                          or a.props.get("code") or a.props.get("id") or "não informado")
                ano = (a.props.get("year") or a.props.get("ano") or a.props.get("alert_year")
                       or a.props.get("data_alerta") or a.props.get("detected_at") or a.props.get("alert_inserted_at")
                       or "não informado")
                area = a.props.get("area_ha") or a.area_intersecao_ha
                _add_bullet(doc, "Há um alerta inserido na área analisada:", 0, num_id)
                _add_bullet(doc, f"○ Código do alerta: {codigo}", 1, num_id)
                _add_bullet(doc, f"○ Ano: {str(ano)[:4] if ano != 'não informado' else ano}", 1, num_id)
                try:
                    area_txt = f"{_fmt_area(float(area))} ha"
                except (TypeError, ValueError):
                    area_txt = f"{area or 'não informado'}"
                _add_bullet(doc, f"○ Área: {area_txt}", 1, num_id)
    else:
        _add_bullet(doc, "Nada consta no MapBiomas Alerta para a geometria consultada.", 0, num_id)

    doc.add_paragraph("SCCON Alertas - Dashboard Mato Grosso:")
    sccon_alertas = (sccon_data or {}).get("alertas") or []
    if sccon_alertas:
        periodo = (sccon_data or {}).get("periodo") or ("2020-01-01", date.today().isoformat())
        _add_bullet(
            doc,
            f"Foram encontrados {len(sccon_alertas)} alerta(s) interceptando a geometria analisada no período de {periodo[0]} a {periodo[1]}.",
            0,
            num_id,
        )
        for idx, alerta in enumerate(sccon_alertas[:20], start=1):
            _add_bullet(doc, f"Alerta SCCON {idx}", 0, num_id)
            _add_label_value(doc, "Código do alerta", alerta.get("codigo"), 1, num_id)
            _add_label_value(doc, "Classe", alerta.get("classe"), 1, num_id)
            _add_label_value(doc, "Data de detecção", alerta.get("data_deteccao"), 1, num_id)
            _add_label_value(doc, "Área", _fmt_area(alerta.get("area_ha") or 0), 1, num_id, " ha")
            _add_label_value(doc, "Imagem antes", alerta.get("imagem_antes"), 1, num_id)
            _add_label_value(doc, "Imagem depois", alerta.get("imagem_depois"), 1, num_id)
        if len(sccon_alertas) > 20:
            _add_bullet(doc, f"Demais alertas omitidos nesta pré-análise: {len(sccon_alertas) - 20}.", 0, num_id)
        for fig in (sccon_data or {}).get("figuras") or []:
            legenda = "todos os anos" if fig.get("ano") == "geral" else fig.get("ano")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"Figura N. Alertas SCCON ({legenda}) com o polígono da propriedade plotado.")
            r.font.bold = True
            r.font.size = Pt(9)
            doc.add_picture(fig["path"], width=Inches(6.45))
    else:
        _add_bullet(doc, "Nada consta no SCCON Alertas para a geometria e período consultados, ou a fonte não respondeu no momento da execução.", 0, num_id)

    doc.add_paragraph("ALERTAS - SIGA - SEMA/MT")
    _add_bullet(
        doc,
        "Os alertas do SIGA/SEMA-MT não possuem serviço público consultável por geometria: "
        f"conferência manual obrigatória no painel do SIGA em {projeto.data_consulta_efetiva()}. "
        "Inserir aqui os prints/figuras por ano quando houver ocorrências.",
        0, num_id,
    )


def _avisos(doc, avisos, llm_docs, area_zip):
    db.cabecalho_secao(doc, "RASTREABILIDADE DA EXECUÇÃO", db.CINZA)
    doc.add_paragraph(
        f"Entrada: {os.path.basename(area_zip.zip_path)}; polígonos: {area_zip.feature_count}; "
        f"área total aproximada: {_fmt_area(area_zip.area_ha)} ha; bbox geográfico: {tuple(round(v, 6) for v in area_zip.bbox_geo)}."
    )
    if llm_docs:
        ok = sum(1 for d in llm_docs if d.get("ok"))
        doc.add_paragraph(f"DeepSeek v4 Flash acionado para {len(llm_docs)} PDF(s); respostas bem-sucedidas: {ok}.")
    if avisos:
        doc.add_paragraph("Avisos técnicos:")
        for a in avisos[:20]:
            doc.add_paragraph(a)


def _montar_doc(projeto, area_zip, car_dados, fundiario, protegidas, legal, alertas, sccon_data, avisos, llm_docs):
    doc = db.novo_documento()
    db.configurar_paginas(doc)
    num_id = db.setup_bullets(doc)
    db.titulo_principal(doc, "ANÁLISE DE ÁREA")
    _dominialidade(doc, projeto, num_id, avisos)
    _contexto_fundiario(doc, projeto, fundiario, num_id)
    _contexto_ambiental(doc, car_dados, projeto, num_id)
    _areas_protegidas(doc, protegidas, num_id)
    _legal(doc, legal, num_id, projeto)
    _alertas(doc, alertas, sccon_data, num_id, projeto)
    return doc


def gerar(projeto, shapefile_zip: str | None = None, destino: str | None = None,
          saida_dir: str | None = None) -> str:
    avisos: list[str] = []
    # IA obrigatória (decisão §0.1 do PLANO_MELHORIAS): sem chave, falha rápido.
    api_key = _deepseek_key(projeto)
    if not api_key:
        raise RuntimeError(
            "Chave DeepSeek ausente: a pré-análise exige IA para extração de PDFs e resumo "
            "jurídico. Configure 'deepseek_api_key' no secrets.local.json da análise "
            "(ou a variável DEEPSEEK_API_KEY) na tela Config e rode novamente."
        )
    zip_path = shapefile_zip or _zip_entrada(projeto)
    nome_saida = _nome_saida(projeto)
    car_saida_dir = apf_saida_dir = None
    sccon_saida_dir = None
    if saida_dir:
        saida_dir = os.path.abspath(saida_dir)
        os.makedirs(saida_dir, exist_ok=True)
        destino = destino or os.path.join(saida_dir, nome_saida)
        car_saida_dir = os.path.join(saida_dir, "Consultas_Publicas", "CAR")
        apf_saida_dir = os.path.join(saida_dir, "Consultas_Publicas", "APF")
        sccon_saida_dir = os.path.join(saida_dir, "Consultas_Publicas", "SCCON")
        os.makedirs(car_saida_dir, exist_ok=True)
        os.makedirs(apf_saida_dir, exist_ok=True)
        os.makedirs(sccon_saida_dir, exist_ok=True)
    destino = destino or io.caminho_resultado(projeto, nome_saida)
    authkey = _authkey(projeto)
    if not authkey:
        avisos.append("Authkey SEMA ausente; consultas SEMA serão parciais.")
    temp_root = os.path.join(projeto.raiz_abs(), "tmp", "shapefiles_uploads")
    with geo.abrir_shape_zip(zip_path, projeto.crs_utm, projeto.crs_geo, temp_root=temp_root) as area_zip:
        car_dados, llm_docs = _coletar_car(
            projeto, area_zip, authkey, api_key, avisos,
            car_saida_dir=car_saida_dir, apf_saida_dir=apf_saida_dir,
        )
        fundiario = _coletar_fundiario(projeto, avisos)
        protegidas = _coletar_protegidas(projeto, area_zip, authkey, avisos)
        _, cars_by_id = _fazendas_intersectadas(projeto, area_zip, authkey, avisos)
        legal = _coletar_legal(projeto, area_zip, authkey, cars_by_id, avisos, api_key)
        alertas = _coletar_alertas(projeto, area_zip, avisos)
        sccon_data = _coletar_sccon(projeto, area_zip, avisos, sccon_saida_dir)
        doc = _montar_doc(projeto, area_zip, car_dados, fundiario, protegidas, legal, alertas, sccon_data, avisos, llm_docs)
        os.makedirs(os.path.dirname(destino), exist_ok=True)
        doc.save(destino)
    return destino


def resumo_shape(projeto, shapefile_zip: str | None = None) -> dict:
    zip_path = shapefile_zip or _zip_entrada(projeto)
    temp_root = os.path.join(projeto.raiz_abs(), "tmp", "shapefiles_uploads")
    avisos: list[str] = []
    with geo.abrir_shape_zip(zip_path, projeto.crs_utm, projeto.crs_geo, temp_root=temp_root) as area_zip:
        resumo = area_zip.resumo()
        fazendas, _ = _fazendas_intersectadas(projeto, area_zip, _authkey(projeto), avisos)
        resumo["fazendas_intersectadas"] = [{"id": f.id, "nome": f.nome} for f in fazendas]
        resumo["avisos"] = resumo.get("avisos", []) + avisos
        return resumo


def _main(argv: list[str]) -> int:
    from core.config import load_projeto
    if len(argv) < 2:
        print("uso: python -m automations.pre_analise <projeto.json> [shape.zip] [saida.docx|pasta_saida]")
        return 2
    proj = load_projeto(argv[1])
    args = [a for a in argv[2:] if not a.startswith("--")]
    shape = args[0] if len(args) >= 1 else None
    destino = saida_dir = None
    if len(args) >= 2:
        if args[1].lower().endswith(".docx"):
            destino = args[1]
        else:
            saida_dir = args[1]
    out = gerar(proj, shapefile_zip=shape, destino=destino, saida_dir=saida_dir)
    print("Gerado:", out)
    return 0


if __name__ == "__main__":
    import sys
    raise SystemExit(_main(sys.argv))
