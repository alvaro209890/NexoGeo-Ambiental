# -*- coding: utf-8 -*-
"""Automação: bloco "Área total" (certificação × matrícula) -> .docx.

Reproduz o antigo ``gerar_area_total.py``, agora dirigido pelo bloco
``area_total.itens`` do projeto (visão de REGISTRO/matrícula). Nenhuma fazenda
escrita no código.

Formato:
    Área total:
      • <nome>: <área certificação> ha (SIGEF/SNCI)
          o Divergência: <área matrícula> ha (Matrícula)

Uso (a partir de software/):
    python -m automations.area_total <projeto.json> [saida.docx]
"""
from __future__ import annotations

from docx.shared import Pt

from core import io, normalize as nz
from core import docx_builder as db

NOME_PADRAO = "Area_Total_Certificacao_Matricula.docx"


def gerar(projeto, destino: str | None = None) -> str:
    """Gera o documento de área total e devolve o caminho salvo."""
    itens = projeto.area_total
    if not itens:
        raise ValueError("projeto sem bloco 'area_total.itens' — nada a gerar.")

    doc = db.novo_documento()
    num_id = db.setup_bullets(doc)

    titulo = doc.add_paragraph()
    titulo.paragraph_format.space_after = Pt(2)
    titulo.add_run("Área total:").font.bold = True

    for it in itens:
        nome = it.get("nome", "")
        tipo = it.get("tipo", "")
        cert = float(it.get("certificacao_ha", 0) or 0)
        matr = float(it.get("matricula_ha", 0) or 0)

        p = doc.add_paragraph()
        db.set_bullet(p, 0, num_id)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
        p.add_run(f"{nome}: ").font.bold = True
        sufixo_tipo = f" ({tipo})" if tipo else ""
        p.add_run(f"{nz.br(cert, 4)} ha{sufixo_tipo}")

        d = doc.add_paragraph()
        db.set_bullet(d, 1, num_id)
        d.paragraph_format.space_after = Pt(0)
        d.add_run("Divergência: ").font.bold = True
        d.add_run(f"{nz.br(matr, nz.casas_decimais(matr))} ha (Matrícula)")

    destino = destino or io.caminho_resultado(projeto, NOME_PADRAO)
    doc.save(destino)
    return destino


def _main(argv: list[str]) -> int:
    from core.config import load_projeto
    if len(argv) < 2:
        print("uso: python -m automations.area_total <projeto.json> [saida.docx]")
        return 2
    proj = load_projeto(argv[1])
    out = gerar(proj, argv[2] if len(argv) > 2 else None)
    print("Gerado:", out)
    return 0


if __name__ == "__main__":
    import sys
    raise SystemExit(_main(sys.argv))
