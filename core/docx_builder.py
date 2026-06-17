# -*- coding: utf-8 -*-
"""core.docx_builder — estilo padrão dos documentos Word (.docx).

Genérico: documento base em Tahoma, numeração multinível (• / o / ▪), bordas e
filetes — extraído dos geradores de contexto (onde `setup_bullets`/`set_bullet`/
`add_bottom_border` estavam copiados em 3 scripts).
"""
from __future__ import annotations

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml

# Cores institucionais
AZUL = RGBColor(0x1F, 0x4E, 0x79)
PRETO = RGBColor(0x00, 0x00, 0x00)
CINZA = RGBColor(0x59, 0x59, 0x59)

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def novo_documento(fonte: str = "Tahoma", tamanho: int = 11) -> Document:
    """Documento novo com o estilo Normal já configurado."""
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = fonte
    normal.font.size = Pt(tamanho)
    return doc


def setup_bullets(doc: Document, num_id: int = 77, abstract_id: int = 77) -> int:
    """Cria a numeração de 3 níveis (• Symbol / o Courier / ▪ Wingdings) e devolve o num_id."""
    numbering = doc.part.numbering_part.element
    abstract = parse_xml(
        f'<w:abstractNum xmlns:w="{_W}" w:abstractNumId="{abstract_id}">'
        '<w:multiLevelType w:val="hybridMultilevel"/>'
        '<w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="bullet"/>'
        '<w:lvlText w:val="&#61623;"/><w:lvlJc w:val="left"/>'
        '<w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>'
        '<w:rPr><w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/></w:rPr></w:lvl>'
        '<w:lvl w:ilvl="1"><w:start w:val="1"/><w:numFmt w:val="bullet"/>'
        '<w:lvlText w:val="o"/><w:lvlJc w:val="left"/>'
        '<w:pPr><w:ind w:left="1440" w:hanging="360"/></w:pPr>'
        '<w:rPr><w:rFonts w:ascii="Courier New" w:hAnsi="Courier New" w:cs="Courier New" w:hint="default"/></w:rPr></w:lvl>'
        '<w:lvl w:ilvl="2"><w:start w:val="1"/><w:numFmt w:val="bullet"/>'
        '<w:lvlText w:val="&#61607;"/><w:lvlJc w:val="left"/>'
        '<w:pPr><w:ind w:left="2160" w:hanging="360"/></w:pPr>'
        '<w:rPr><w:rFonts w:ascii="Wingdings" w:hAnsi="Wingdings" w:hint="default"/></w:rPr></w:lvl>'
        '</w:abstractNum>'
    )
    num = parse_xml(
        f'<w:num xmlns:w="{_W}" w:numId="{num_id}"><w:abstractNumId w:val="{abstract_id}"/></w:num>'
    )
    numbering.insert(0, abstract)
    numbering.append(num)
    return num_id


def set_bullet(paragraph, ilvl: int, num_id: int):
    """Aplica o nível de bullet a um parágrafo."""
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    el_ilvl = OxmlElement("w:ilvl")
    el_ilvl.set(qn("w:val"), str(ilvl))
    el_num = OxmlElement("w:numId")
    el_num.set(qn("w:val"), str(num_id))
    numPr.append(el_ilvl)
    numPr.append(el_num)
    pPr.insert(0, numPr)


def add_bottom_border(paragraph, color: str = "1F4E79", sz: str = "14"):
    """Filete inferior (usado nos cabeçalhos de seção)."""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.insert(0, pbdr)


def configurar_paginas(doc: Document):
    """Margens proximas ao modelo de analise."""
    for sec in doc.sections:
        sec.top_margin = Inches(0.85)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(0.85)
        sec.right_margin = Inches(0.75)


def titulo_principal(doc: Document, texto: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(texto)
    r.font.name = "Tahoma"
    r.font.bold = True
    r.font.size = Pt(15)
    return p


def cabecalho_secao(doc: Document, texto: str, cor: RGBColor = PRETO):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(texto)
    r.font.name = "Tahoma"
    r.font.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = cor
    add_bottom_border(p)
    return p


def paragrafo_fonte(doc: Document, texto: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(texto)
    r.font.size = Pt(8)
    r.font.italic = True
    r.font.color.rgb = CINZA
    return p
